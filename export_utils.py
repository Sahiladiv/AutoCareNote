from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def export_pdf(soap, summary, path="output.pdf"):
    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    flow = []
    flow.append(Paragraph("SOAP Note", styles['Heading1']))
    flow.append(Paragraph(soap, styles['Normal']))
    flow.append(Spacer(1, 20))
    flow.append(Paragraph("Patient Summary", styles['Heading1']))
    flow.append(Paragraph(summary, styles['Normal']))
    doc.build(flow)

def export_docx(soap, summary, path="output.docx"):
    doc = Document()
    doc.add_heading("SOAP Note", 0)
    doc.add_paragraph(soap)
    doc.add_heading("Patient Summary", 0)
    doc.add_paragraph(summary)
    doc.save(path)
